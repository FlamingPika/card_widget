from os import listdir
from os.path import isfile, join
from docx import Document
"""
Python script to pack all the file names into a single list.

1. Copy all the files into /src folder.
2. Change the condition of file type (.jpeg, etc")
3. Run the code and the file names should be in the terminal

"""

def generate_list():
    # onlyfiles = ["src/"+f for f in listdir("src") if isfile(join("src", f)) and (f.endswith('.jpeg') or f.endswith('.jpg') or f.endswith('.png') or f.endswith('.gif'))]
    onlyfiles = ["src/"+f for f in listdir("src") if isfile(join("src", f)) and f.endswith('.gif')]

    # Number of files
    print(f"Total number of files is {len(onlyfiles)}")

    raw_content = [f[:-4].split("_") for f in listdir("src") if isfile(join("src", f)) and f.endswith('.gif')]
    content = []

    for i in raw_content:
        content.append(i[1].capitalize() + " of "+i[0].capitalize())

    print(content)

"""
@desc:
A function to generate the description of tarot cards with the format that is compatible with JS file format inside the code

@param: 
- path = the path to where the file is saved
- filename = the filename that contains the description
- extension = the file extension

@return (printed):
A dictionary with the path to the image as the key, and [title, content] as the value
"""
def generate_desc(path, filename, extension):
    number = {'one': '1', 'two': '2', 'three': '3', 'four': '4', 'five': '5', 'six': '6', 'seven': '7', 'eight': '8', 'nine': '9', 'ten': '10',}
    doc = Document(join('src', filename))
    content = {}
    line = 0
    title = ""
    while line < len(doc.paragraphs):
        if doc.paragraphs[line].text == '': 
            raw_title = doc.paragraphs[line + 1].text
            title = raw_title.lower().split()
            if title[0] in number:
                title = join(path, title[2] + '_' + number[title[0]] + extension)
            else:
                title = join(path, title[2] + '_' + title[0] + extension)
            content[title] = [raw_title,""]
            line += 2
        else:
            content[title][1] = doc.paragraphs[line].text
            line += 1


    print(f"Total number of files is {len(content)}")

    image_list = ['src/discs_daughter.gif', 'src/cups_8.gif', 'src/wands_daughter.gif', 'src/swords_9.gif', 'src/wands_9.gif', 'src/wands_son.gif', 'src/wands_8.gif', 'src/swords_8.gif', 'src/cups_9.gif', 'src/swords_shaman.gif', 'src/discs_9.gif', 'src/wands_10.gif', 'src/discs_priestess.gif', 'src/discs_8.gif', 'src/cups_10.gif', 'src/cups_priestess.gif', 'src/wands_ace.gif', 'src/swords_priestess.gif', 'src/wands_priestess.gif', 'src/discs_3.gif', 'src/cups_shaman.gif', 'src/cups_ace.gif', 'src/discs_2.gif', 'src/swords_daughter.gif', 'src/discs_6.gif', 'src/wands_shaman.gif', 'src/swords_ace.gif', 'src/discs_7.gif', 'src/discs_5.gif', 'src/discs_son.gif', 'src/discs_10.gif', 'src/discs_4.gif', 'src/cups_7.gif', 'src/cups_son.gif', 'src/swords_6.gif', 'src/swords_10.gif', 'src/wands_6.gif', 'src/wands_7.gif', 'src/cups_daughter.gif', 'src/swords_7.gif', 'src/cups_6.gif', 'src/cups_4.gif', 'src/swords_5.gif', 'src/wands_5.gif', 'src/wands_4.gif', 'src/swords_4.gif', 'src/cups_5.gif', 'src/discs_shaman.gif', 'src/discs_ace.gif', 'src/cups_2.gif', 'src/swords_3.gif', 'src/wands_3.gif', 'src/wands_2.gif', 'src/swords_son.gif', 'src/swords_2.gif', 'src/cups_3.gif'];

    print(content)


generate_desc('src', 'TarotDescriptions.docx', '.gif')


